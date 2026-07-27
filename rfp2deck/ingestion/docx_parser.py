from __future__ import annotations

from dataclasses import dataclass
from io import BytesIO
from pathlib import Path
from typing import Union

from docx import Document

from rfp2deck.core.logging import get_logger

log = get_logger(__name__)


@dataclass
class ParsedDoc:
    text: str
    paragraph_count: int
    table_count: int = 0


def parse_docx(path_or_bytes: Union[Path, bytes]) -> ParsedDoc:
    source = "bytes" if isinstance(path_or_bytes, bytes) else str(path_or_bytes)
    try:
        if isinstance(path_or_bytes, bytes):
            d = Document(BytesIO(path_or_bytes))
        else:
            d = Document(path_or_bytes)
    except Exception:
        log.exception("Failed to open DOCX (source=%s)", source)
        raise
    lines = []
    paragraph_count = 0
    for paragraph_number, paragraph in enumerate(d.paragraphs, start=1):
        text = paragraph.text.strip() if paragraph.text else ""
        if not text:
            continue
        paragraph_count += 1
        style_name = getattr(getattr(paragraph, "style", None), "name", "") or ""
        lines.append(
            f'[PARAGRAPH {paragraph_number}][STYLE "{style_name}"] {text}'
        )

    for table_number, table in enumerate(d.tables, start=1):
        lines.append(f"\n--- TABLE {table_number} ---")
        for row_number, row in enumerate(table.rows, start=1):
            cells = [" ".join(cell.text.split()) for cell in row.cells]
            lines.append(
                f"[TABLE {table_number}][ROW {row_number}] "
                + " | ".join(f"Column {index + 1}: {value}" for index, value in enumerate(cells))
            )

    parsed = ParsedDoc(
        text="\n".join(lines),
        paragraph_count=paragraph_count,
        table_count=len(d.tables),
    )
    log.info(
        "Parsed DOCX (source=%s): %d paragraphs, %d tables, %d chars",
        source,
        parsed.paragraph_count,
        parsed.table_count,
        len(parsed.text),
    )
    return parsed
